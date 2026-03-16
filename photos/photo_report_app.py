import os
import io
import tempfile
from datetime import datetime

from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QLabel,
    QFileDialog,
    QListWidget,
    QListWidgetItem,
    QMessageBox,
    QGroupBox,
    QProgressBar,
)
from PyQt5.QtCore import Qt

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Twips, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image

import shared.global_data as global_data


# Template'teki satir yüksekligi (twips cinsinden)
ROW_HEIGHT_TWIPS = 3768
CELL_WIDTH_MM = 86
PHOTO_WIDTH_MM = 80
# Sikistirma ayarlari
COMPRESS_MAX_WIDTH = 1200
COMPRESS_QUALITY = 75


def compress_photo(photo_path):
    """Fotoğrafı sıkıştırıp geçici dosya olarak döndürür."""
    img = Image.open(photo_path)

    # EXIF orientation düzelt
    try:
        from PIL import ExifTags
        for orientation_key in ExifTags.TAGS:
            if ExifTags.TAGS[orientation_key] == "Orientation":
                break
        exif = img._getexif()
        if exif and orientation_key in exif:
            orient = exif[orientation_key]
            if orient == 3:
                img = img.rotate(180, expand=True)
            elif orient == 6:
                img = img.rotate(270, expand=True)
            elif orient == 8:
                img = img.rotate(90, expand=True)
    except Exception:
        pass

    # Boyut küçült
    if img.width > COMPRESS_MAX_WIDTH:
        ratio = COMPRESS_MAX_WIDTH / img.width
        new_h = int(img.height * ratio)
        img = img.resize((COMPRESS_MAX_WIDTH, new_h), Image.LANCZOS)

    # RGB'ye çevir (PNG transparanlık sorunu)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=COMPRESS_QUALITY, optimize=True)
    buf.seek(0)
    return buf


def _set_row_height(row, height_twips):
    """Satır yüksekliğini twips cinsinden ayarlar."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height_twips))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


class PhotoReportApp(QMainWindow):
    IMAGE_FILTER = "Fotoğraflar (*.jpg *.jpeg *.png)"
    CATEGORIES = {
        "PRE": {
            "title": "Pre",
            "template": "PRE.docx",
            "color": "#03A9F4",
        },
        "POST": {
            "title": "Post",
            "template": "POST.docx",
            "color": "#4CAF50",
        },
        "TEARDOWN": {
            "title": "Teardown",
            "template": "TEARDOWN.docx",
            "color": "#FF9800",
        },
        "HANDLE_SIDE_COVER": {
            "title": "Handle Side Cover",
            "template": "HANDLE_SIDE_COVER.docx",
            "color": "#9C27B0",
        },
    }

    # Her kategorinin başlık metni (template'teki ile aynı)
    CATEGORY_TITLES = {
        "PRE": "PRE TEST PHOTOS",
        "POST": "POST TEST PHOTOS",
        "TEARDOWN": "TEAR DOWN PHOTOS",
        "HANDLE_SIDE_COVER": "Handle Side Cover Photos",
    }

    def __init__(self, main_window=None):
        super().__init__()
        self.main_window = main_window
        self.setWindowTitle("Photo Report")
        self.resize(900, 760)

        self.selected_files = {key: [] for key in self.CATEGORIES}
        self.list_widgets = {}

        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QVBoxLayout(main_widget)

        title = QLabel("Photo Report Modülü")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("font-size: 20px; font-weight: bold; margin-bottom: 10px;")
        main_layout.addWidget(title)

        subtitle = QLabel("Her kategori için çoklu fotoğraf seçebilir ve ayrı Word raporu oluşturabilirsiniz.")
        subtitle.setAlignment(Qt.AlignCenter)
        subtitle.setStyleSheet("color: #616161; margin-bottom: 8px;")
        main_layout.addWidget(subtitle)

        for category, cfg in self.CATEGORIES.items():
            main_layout.addWidget(self._build_category_section(category, cfg))

        action_row = QHBoxLayout()

        self.btn_generate = QPushButton("Seçilen Kategoriler İçin Raporları Oluştur")
        self.btn_generate.setStyleSheet(
            "font-size: 15px; padding: 12px; background-color: #1976D2; color: white; font-weight: bold;"
        )
        self.btn_generate.clicked.connect(self.generate_reports)

        self.btn_back = QPushButton("Ana Menüye Dön")
        self.btn_back.setStyleSheet("font-size: 15px; padding: 12px; background-color: #9E9E9E; color: white;")
        self.btn_back.clicked.connect(self.close_and_return)

        action_row.addWidget(self.btn_generate)
        action_row.addWidget(self.btn_back)

        main_layout.addLayout(action_row)

        self.progress = QProgressBar()
        self.progress.setMinimum(0)
        self.progress.setValue(0)
        self.progress.setFormat("Hazır")
        main_layout.addWidget(self.progress)

        self.lbl_output = QLabel("Çıktı klasörü: tempfiles/photo_reports")
        self.lbl_output.setStyleSheet("color: #616161; margin-top: 4px;")
        main_layout.addWidget(self.lbl_output)

    def _build_category_section(self, category, cfg):
        group = QGroupBox(cfg["title"])
        group.setStyleSheet(f"QGroupBox {{ font-weight: bold; color: {cfg['color']}; }}")

        layout = QVBoxLayout(group)

        button_row = QHBoxLayout()

        btn_select = QPushButton(f"{cfg['title']} fotoğraf seç")
        btn_select.clicked.connect(lambda _, c=category: self.select_photos(c))
        button_row.addWidget(btn_select)

        btn_up = QPushButton("▲")
        btn_up.setFixedWidth(36)
        btn_up.clicked.connect(lambda _, c=category: self.move_photo(c, -1))
        button_row.addWidget(btn_up)

        btn_down = QPushButton("▼")
        btn_down.setFixedWidth(36)
        btn_down.clicked.connect(lambda _, c=category: self.move_photo(c, 1))
        button_row.addWidget(btn_down)

        btn_reset = QPushButton("Tümünü temizle")
        btn_reset.clicked.connect(lambda _, c=category: self.clear_category(c))
        button_row.addWidget(btn_reset)

        layout.addLayout(button_row)

        list_widget = QListWidget()
        list_widget.setSelectionMode(QListWidget.ExtendedSelection)
        list_widget.setMinimumHeight(80)
        layout.addWidget(list_widget)

        self.list_widgets[category] = list_widget
        self._refresh_list(category)

        return group

    def select_photos(self, category):
        files, _ = QFileDialog.getOpenFileNames(
            self,
            f"{self.CATEGORIES[category]['title']} fotoğraflarını seç",
            "",
            self.IMAGE_FILTER,
        )

        if not files:
            return

        valid_ext = {".jpg", ".jpeg", ".png"}
        dedup = set(self.selected_files[category])

        for path in files:
            ext = os.path.splitext(path)[1].lower()
            if ext in valid_ext and path not in dedup:
                self.selected_files[category].append(path)
                dedup.add(path)

        self._refresh_list(category)

    def move_photo(self, category, direction):
        list_widget = self.list_widgets[category]
        selected = list_widget.selectedIndexes()
        if not selected or not self.selected_files[category]:
            return

        row = selected[0].row()
        files = self.selected_files[category]
        new_row = row + direction

        if new_row < 0 or new_row >= len(files):
            return

        files[row], files[new_row] = files[new_row], files[row]
        self._refresh_list(category)
        list_widget.setCurrentRow(new_row)

    def clear_category(self, category):
        self.selected_files[category] = []
        self._refresh_list(category)

    def _refresh_list(self, category):
        list_widget = self.list_widgets[category]
        list_widget.clear()

        files = self.selected_files[category]
        if not files:
            list_widget.addItem(QListWidgetItem("Henüz fotoğraf seçilmedi"))
            list_widget.item(0).setFlags(Qt.NoItemFlags)
            return

        for i, photo in enumerate(files):
            list_widget.addItem(QListWidgetItem(f"{i+1}. {os.path.basename(photo)}"))

    def _ensure_output_dir(self):
        root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        test_no = str(global_data.config.get("TEST_NO") or "UNSPECIFIED").replace("/", "_").replace("\\", "_").strip()
        date_part = datetime.now().strftime("%Y%m%d")
        out_dir = os.path.join(root_dir, "tempfiles", "photo_reports", test_no, date_part)
        os.makedirs(out_dir, exist_ok=True)
        return out_dir

    def _resolve_template_path(self, category):
        base_dir = os.path.dirname(os.path.abspath(__file__))
        template_name = self.CATEGORIES[category]["template"]
        return os.path.join(base_dir, "templates", template_name)

    def _safe_output_name(self, category, output_dir):
        test_no = global_data.config.get("TEST_NO") or "UNSPECIFIED"
        normalized_test = str(test_no).replace("/", "_").replace("\\", "_").strip()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = f"photo_report_{category.lower()}_{normalized_test}_{timestamp}"

        candidate = os.path.join(output_dir, f"{base_name}.docx")
        idx = 1
        while os.path.exists(candidate):
            candidate = os.path.join(output_dir, f"{base_name}_{idx}.docx")
            idx += 1

        return candidate

    def _chunk_photos(self, photos, size=6):
        return [photos[idx : idx + size] for idx in range(0, len(photos), size)]

    def _get_title_paragraph(self, doc):
        """Template'teki başlık paragrafını bulur."""
        for p in doc.paragraphs:
            if p.text.strip():
                return p
        return None

    def _copy_paragraph_format(self, source_para):
        """Paragraf formatını (font, alignment, bold vb.) dict olarak saklar."""
        fmt = {
            "alignment": source_para.alignment,
            "bold": None,
            "font_size": None,
            "font_name": None,
        }
        if source_para.runs:
            run = source_para.runs[0]
            fmt["bold"] = run.bold
            fmt["font_size"] = run.font.size
            fmt["font_name"] = run.font.name
        return fmt

    def _insert_title(self, doc, title_text, fmt):
        """Başlık paragrafı ekler (template formatında)."""
        para = doc.add_paragraph()
        para.alignment = fmt.get("alignment") or WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title_text)
        run.bold = fmt.get("bold", True)
        if fmt.get("font_size"):
            run.font.size = fmt["font_size"]
        if fmt.get("font_name"):
            run.font.name = fmt["font_name"]
        return para

    def _insert_photo_into_cell(self, cell, photo_buf):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = paragraph.add_run()
        run.add_picture(photo_buf, width=Mm(PHOTO_WIDTH_MM))

    def _add_photo_table(self, doc, photo_chunk, compressed_cache):
        """3 satır x 2 sütun tablo ekler, template boyutlarında."""
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"

        # Satir yüksekliklerini template ile ayni yap
        for row in table.rows:
            _set_row_height(row, ROW_HEIGHT_TWIPS)

        for idx in range(6):
            row_idx = idx // 2
            col_idx = idx % 2
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            if idx < len(photo_chunk):
                photo_path = photo_chunk[idx]
                # Sıkıştırılmış halini al (cache'ten veya yeni oluştur)
                if photo_path not in compressed_cache:
                    compressed_cache[photo_path] = compress_photo(photo_path)
                buf = compressed_cache[photo_path]
                buf.seek(0)
                self._insert_photo_into_cell(cell, buf)

    def _build_document_from_template(self, template_path, photos, category):
        doc = Document(template_path)

        # Başlık formatını kaydet
        title_para = self._get_title_paragraph(doc)
        title_text = self.CATEGORY_TITLES.get(category, "PHOTOS")
        title_fmt = self._copy_paragraph_format(title_para) if title_para else {"bold": True}

        # Template'teki tabloları ve boş paragrafları kaldır
        # (İlk sayfa: başlık paragrafı + boş tablo)
        for table in list(doc.tables):
            table._element.getparent().remove(table._element)

        # Boş paragrafları temizle (başlık hariç)
        body = doc.element.body
        for p_elem in list(body.findall(qn("w:p"))):
            text = p_elem.text or ""
            # Run'lardaki text'i de topla
            for r in p_elem.findall(qn("w:r")):
                t = r.find(qn("w:t"))
                if t is not None and t.text:
                    text += t.text
            if not text.strip():
                body.remove(p_elem)

        photo_chunks = self._chunk_photos(photos, size=6)
        compressed_cache = {}

        for page_no, chunk in enumerate(photo_chunks, start=1):
            if page_no > 1:
                # Yeni sayfa + başlık
                doc.add_page_break()
                self._insert_title(doc, title_text, title_fmt)

            self._add_photo_table(doc, chunk, compressed_cache)

        return doc

    def _render_category_report(self, category, output_path):
        template_path = self._resolve_template_path(category)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template bulunamadı: {template_path}")

        doc = self._build_document_from_template(
            template_path, self.selected_files[category], category
        )
        doc.save(output_path)

    def generate_reports(self):
        selected_categories = [key for key, files in self.selected_files.items() if files]
        if not selected_categories:
            QMessageBox.warning(self, "Uyarı", "Rapor üretmek için en az bir kategoriye fotoğraf ekleyin.")
            return

        output_dir = self._ensure_output_dir()

        total_photos = sum(len(self.selected_files[c]) for c in selected_categories)
        self.progress.setMinimum(0)
        self.progress.setMaximum(total_photos)
        self.progress.setValue(0)

        created_files = []
        progress_count = 0

        try:
            for category in selected_categories:
                self.progress.setFormat(f"{self.CATEGORIES[category]['title']} hazırlanıyor...")
                QApplication.processEvents()

                output_path = self._safe_output_name(category, output_dir)
                self._render_category_report(category, output_path)
                created_files.append(output_path)

                progress_count += len(self.selected_files[category])
                self.progress.setValue(progress_count)
                QApplication.processEvents()

            result_text = "\n".join(created_files)
            QMessageBox.information(
                self,
                "Başarılı",
                f"{len(created_files)} rapor başarıyla oluşturuldu.\n\nÇıktılar:\n{result_text}",
            )
            self.progress.setFormat("Tamamlandı")
        except Exception as exc:
            QMessageBox.critical(self, "Hata", f"Rapor oluşturulurken hata oluştu:\n{exc}")
            self.progress.setFormat("Hata!")

    def closeEvent(self, event):
        if self.main_window is not None:
            self.main_window.show()
        event.accept()

    def close_and_return(self):
        self.close()
